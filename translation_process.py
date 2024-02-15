import openai
from langdetect import detect, DetectorFactory
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv
import spacy, re, os, time, fitz

def split_into_paragraphs(text):
    return [paragraph for paragraph in text.split('\n') if paragraph.strip() != '']

def get_docx_paragraphs(file_path):
    doc = Document(file_path)
    whole_text = ""
    for paragraph in doc.paragraphs:
        whole_text += paragraph.text + "\n"
    return whole_text

def get_pdf_paragraphs(file_path):
    # Open the PDF file
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        # Extract text from the page
        page_text = page.get_text("text")
        paragraphs = page_text.split('\n')
        for paragraph in paragraphs:
            text += paragraph + "\n"
    doc.close()
    return text

def remove_footnote_markers(text):
    pattern = r'\[\d+\]|\d+\s?|\(\d+\)' 
    cleaned_text = re.sub(pattern, '', text)
    return cleaned_text


def identify_entities(text, source_language, target_language):
    nlp = ner_models[source_language]
    doc = nlp(text)
    entities = {}

    for ent in doc.ents:
        if ent.text not in entities and ent.label_ in ['PERSON', 'GPE', 'LOC', 'ORG', 'FAC', 'EVENT', 'NORP', 'WORK_OF_ART', 'PRODUCT']:
            full_sentence = ent.sent.text
            entities[ent.text] = translate_entities_context(ent.text, full_sentence, source_language=source_language, target_language=target_language, client=client)
            print(f"{ent.text} -> {entities[ent.text]}")
    return entities

def split_into_sentences(paragraph):
    return paragraph.split('. ')

def translate_entities_context(entity, context_window, source_language, target_language, client):
    print(f"translating {entity}")
    context_messages = [{"role": "system", "content": f"Translate this text '{entity}' from {source_language} to {target_language} academically, given this context: '{context_window}' translate this text '{entity}'"}]

    while True:
        try:
            response = client.chat.completions.create(model="gpt-4",
                                                      messages=context_messages,
                                                      temperature=0.4,
                                                      max_tokens=6969)
            return response.choices[0].message.content.strip()

        except openai.RateLimitError:
            time.sleep(65)
        except openai.OpenAIError as e:
            raise e
        
def translate_text(sentences, prompt, client):
    context_messages = [{"role": "system", "content": prompt}]
    for sentence in sentences[-3:]:  # Keep the context of the last 3 sentences plus the current one
        context_messages.append({"role": "user", "content": sentence})

    while True:
        try:
            response = client.chat.completions.create(model="gpt-4",
                                                      messages=context_messages,
                                                      temperature=0.4,
                                                      max_tokens=6969)
            return response.choices[0].message.content.strip()

        except openai.RateLimitError:
            time.sleep(65)
        except openai.OpenAIError as e:
            raise e


def replace_entities(text, mapping):
    regex_pattern = '|'.join(re.escape(key) for key in mapping.keys())
    def replace_match(match):
        return mapping[match.group(0)]
    modified_text = re.sub(regex_pattern, replace_match, text)
    return modified_text


def create_translated_docx(paragraphs, prompt, client, output_file_path):
    translated_doc = Document()
    curr = 0
    for paragraph in paragraphs:
        sentences = split_into_sentences(paragraph)
        translated_paragraph = []
        for i in range(len(sentences)):
            context_sentences = sentences[max(0, i-3):i+1] 
            translated_sentence = translate_text(context_sentences, prompt, client)
            translated_paragraph.append(translated_sentence.split('\n')[-1]) 
        print(f"{translated_paragraph} ix-> {curr}")
        translated_doc.save(output_file_path)
        curr += 1
        p = translated_doc.add_paragraph(' '.join(translated_paragraph))
    translated_doc.save(output_file_path)


if (__name__ == '__main__'):

    # This is specific to .docx :(
    
    # Load target lang, input/output file, api key, client -------------------------------------------------
    load_dotenv('config.env')
    API_KEY = os.getenv('API_KEY')
    client = OpenAI(api_key=API_KEY)

    file_path = '/Users/williamzhao/translation-engine/input_file/Dynastic Histories and Kinship Business(1).docx'
    output_file_path = 'output_files/refined Dynastic Histories and Kinship Business(1).docx'
    target_language = 'en'

    # Selection of Correct NER Model -------------------------------------------------
    # DetectorFactory.seed = 0
    # entire_text = get_docx_paragraphs(file_path=file_path)
    # print(entire_text)
    # source_language = detect(entire_text)
    
    # print(f"source language {source_language}")
    # ner_models = {}

    # if source_language == "en":
    #     ner_models['en'] = spacy.load('en' + '_core_web_sm')
    # if (source_language == "zh-cn"):
    #     source_language = "zh"
    #     ner_models[source_language] = spacy.load(f'{source_language}_core_web_sm')
    # if (source_language == "ru"):
    #     ner_models[source_language] = spacy.load(f'{source_language}_core_web_sm')


    # # Translation Process -------------------------------------------------

    # translation_prompt = f"Ignore text not in {source_language}. Translate from {source_language} to {target_language}, ensuring academic quality, correct punctuation, and complete sentences."
    # entity_mapping = identify_entities(entire_text, source_language=source_language, target_language=target_language)
    # print(entity_mapping)
    # modified_text_with_entities = replace_entities(entire_text, entity_mapping)
    # text_without_footnote = remove_footnote_markers(modified_text_with_entities)
    # paragraphs = split_into_paragraphs(text_without_footnote)
    # create_translated_docx(paragraphs, translation_prompt, client, output_file_path)

    # Translation Process -------------------------------------------------


    # Refining Process ---------------------------------------------------
    refinement_prompt = f"Review and Correct the {target_language} text, return the refined text and no OTHER comments: Please carefully analyze the translated text provided below. Identify and correct any errors related to punctuation, grammar, and sentence structure. Ensure that all sentences are complete and logically coherent. Pay close attention to the proper use of commas, periods, and other punctuation marks. Adjust any awkward or unclear phrasing to improve readability and fluency. Your goal is to refine the text to a high standard of language quality, ensuring it is clear, correct, and professionally presented."
    paragraphs = split_into_paragraphs(get_docx_paragraphs(output_file_path))
    # print(paragraphs)
    create_translated_docx(paragraphs, refinement_prompt, client, output_file_path)

    



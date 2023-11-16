# Helper functions for working with .docx files
from docx import Document
from .translation import translate

def get_paragraphs(file_path):
    doc = Document(file_path)
    return [(paragraph.text, paragraph.style) for paragraph in doc.paragraphs if paragraph.text.strip() != '']

def translate_docx(document) -> list:
    '''
    The following function accepts a .docx file
    It extracts the text from the file
    It splits the text into paragraphs
    It calls the translate function on each paragraph
    The function returns a list of English paragraphs
    '''
    # The list of English paragraphs
    english = []

    # Extract the text from the file
    text = get_paragraphs(document)

    # Iterate through each paragraph
    for paragraph in text:
        # Translate the paragraph
        english.append(translate(paragraph[0]))

    # Return the list of English paragraphs
    return english

def write_docx(english) -> None:
    '''
    The following function accepts a list of English paragraphs
    It writes the paragraphs to a .docx file
    The function returns nothing
    '''

    # Create a new document
    document = Document()

    # Iterate through each paragraph
    for paragraph in english:
        # Add the paragraph to the document
        document.add_paragraph(paragraph)

    # Save the document
    document.save('translated.docx')

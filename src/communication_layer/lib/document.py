# Helper functions for working with .docx files
from docx import Document
from .translation import translate

'''
The following function accepts a .docx file
It extracts the text from the file
It splits the text into paragraphs
It calls the translate function on each paragraph
The function returns a list of English paragraphs
'''

def get_paragraphs(file_path):
    doc = Document(file_path)
    return [(paragraph.text, paragraph.style) for paragraph in doc.paragraphs if paragraph.text.strip() != '']

def translate_docx(document) -> list:
    # The list of English paragraphs
    english = []

    # Extract the text from the file
    text = get_paragraphs(document)

    # Iterate through each paragraph
    for paragraph in text:
        # Translate the paragraph
        english.append(translate(paragraph[0]))

    # Return the list of English paragraphs
    return english

'''
The following function accepts a .docx file
as well as a list of English paragraphs
It writes the paragraphs to the file
The function returns nothing
'''
def write_docx(docx, english) -> None:
    # The document to be written
    document = Document()

    # Iterate through each paragraph
    for paragraph in english:
        # Write the paragraph to the document
        document.add_paragraph(paragraph)

    # Save the document
    document.save(docx)



from openai import OpenAI
from docx import Document
from dotenv import load_dotenv
import spacy, re, os, openai, time

# Load the Chinese language model for spaCy
nlp = spacy.load("zh_core_web_sm")

class EntityRecognizer:
    def __init__(self, nlp_model):
        self.nlp = nlp_model

    def identify_entities(self, text, source_language, target_language, client):
        doc = self.nlp(text)
        entities = {}
        for ent in doc.ents:
            if ent.text not in entities and ent.label_ in ['PERSON', 'GPE', 'LOC', 'ORG']:
                entities[ent.text] = self.translate_text_without_context(ent.text, source_language, target_language, client)
        return entities

    def translate_text_without_context(self, sentence, source_language, target_language, client):
        # Context message for translation
        context_message = {"role": "system", "content": f"Translate from {source_language} to {target_language}."}
        completion_message = {"role": "user", "content": sentence}
    
        while True:
            try:
                response = client.chat.completions.create(
                    model="gpt-4",
                    messages=[context_message, completion_message],
                    temperature=0,
                    max_tokens=6969
                )
                return response.choices[0].message.content.strip()

            except openai.RateLimitError:
                time.sleep(65)
            except openai.OpenAIError as e:
                raise e

class DocumentManager:
    @staticmethod
    def split_into_paragraphs(file_path):
        doc = Document(file_path)
        return [(paragraph.text, paragraph.style) for paragraph in doc.paragraphs if paragraph.text.strip() != '']

    @staticmethod
    def get_entire_text(file_path):
        doc = Document(file_path)
        whole_text = ""
        for paragraph in doc.paragraphs:
            whole_text += paragraph.text + "\n"
        return whole_text

    @staticmethod
    def split_into_sentences(paragraph):
        return paragraph.split('. ')

class TranslationEngine:
    @staticmethod
    def replace_entities(text, mapping):
        regex_pattern = '|'.join(re.escape(key) for key in mapping.keys())
        def replace_match(match):
            return mapping[match.group(0)]
        return re.sub(regex_pattern, replace_match, text)

    def translate_text(self, sentences, source_language, target_language, client):
        sentences = self.replace_entities(sentences, entity_mapping)
        context_messages = [{"role": "system", "content": f"Translate from {source_language} to {target_language} and IGNORE text in {target_language}.."}]
        for sentence in sentences[-4:]:  # Keep the context of the last 3 sentences plus the current one
            context_messages.append({"role": "user", "content": sentence})

        while True:
            try:
                response = client.chat.completions.create(model="gpt-4",
                                                          messages=context_messages,
                                                          temperature=0,
                                                          max_tokens=6969)
                return response.choices[0].message.content.strip()

            except openai.RateLimitError:
                time.sleep(65)
            except openai.OpenAIError as e:
                raise e

class Translator:
    def __init__(self, source_language, target_language, client):
        self.source_language = source_language
        self.target_language = target_language
        self.client = client
        self.entity_recognizer = EntityRecognizer(nlp)
        self.translation_engine = TranslationEngine()
        self.document_manager = DocumentManager()

    def create_translated_document(self, file_path, output_file_path):
        paragraphs = self.document_manager.split_into_paragraphs(file_path)
        translated_doc = Document()
        curr = 0
        for paragraph, style in paragraphs:
            sentences = self.document_manager.split_into_sentences(paragraph)
            translated_paragraph = []
            for i in range(len(sentences)):
                context_sentences = sentences[max(0, i-3):i+1] 
                translated_sentence = self.translation_engine.translate_text(context_sentences, self.source_language, self.target_language, self.client)
                translated_paragraph.append(translated_sentence.split('\n')[-1]) 
            print(f"{translated_paragraph} ix-> {curr}")
            translated_doc.save(output_file_path)
            curr += 1
            p = translated_doc.add_paragraph(' '.join(translated_paragraph))
            p.style = style
        translated_doc.save(output_file_path)

if __name__ == '__main__':

    load_dotenv('config.env')
    API_KEY = os.getenv('API_KEY')
    client = OpenAI(api_key=API_KEY)

    file_path = '/path/to/your/input/file.docx'
    output_file_path = '/path/to/your/output/file-translated.docx'
    source_language = 'your source language'
    target_language = 'your target language'

    # Create an instance of the Translator class
    translator = Translator(source_language, target_language, client)

    # Extract the entire text of the document for entity recognition
    entire_text = DocumentManager.get_entire_text(file_path)

    # Identify entities in the text
    entity_mapping = translator.entity_recognizer.identify_entities(entire_text, source_language, target_language, client)

    # Translate the document and save the result
    translator.create_translated_document(file_path, output_file_path)

    print(f"Translation completed. The translated document is saved at {output_file_path}")



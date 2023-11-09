import re
import time
import openai
from openai import OpenAI

# client = OpenAI()
from docx import Document
import inflect
from dotenv import load_dotenv
import os

# Load the API key from an .env file
load_dotenv('config.env')
API_KEY = os.getenv('API_KEY')
client = OpenAI(api_key=API_KEY)

# Initialize inflect engine
p = inflect.engine()

def calculate_token_length(messages):
    # This function calculates the total token length for a list of messages
    return sum(len(message['content'].split()) for message in messages)

def reduce_context_size(context_messages, max_tokens):
    # This function reduces the size of context_messages to be within max_tokens
    while calculate_token_length(context_messages) > max_tokens:
        # Remove the oldest message
        context_messages.pop(0)
    return context_messages

def handle_rate_limited_request(context_messages, max_tokens=4000):
    max_context_tokens = 8192 - max_tokens
    while True:
        try:
            # Make sure the context_messages don't exceed the token limit.
            if calculate_token_length(context_messages) > max_context_tokens:
                # Reduce the size of context_messages to fit within the limit
                context_messages = reduce_context_size(context_messages, max_context_tokens)

            response = client.chat.completions.create(model="gpt-4",
            messages=context_messages,
            temperature=0,
            max_tokens=max_tokens)
            return response
        except openai.RateLimitError:
            time.sleep(65)
        except openai.OpenAIError as e:
            raise

def segment_text(text, max_segment_size=800):
    """
    Splits the text into segments, each having a length less than the max_segment_size.
    The max_segment_size should account for context and translation instructions.
    """
    sentences = re.split(r'(?<=[.!?]) +', text)
    segments = []
    current_segment = ''
    current_size = 0
    
    for sentence in sentences:
        sentence_size = len(sentence)
        # Check if adding the next sentence would exceed the max_segment_size
        if current_size + sentence_size > max_segment_size:
            segments.append(current_segment.strip())
            current_segment = sentence + ' '
            current_size = sentence_size
        else:
            current_segment += sentence + ' '
            current_size += sentence_size
    
    # Add the last segment if it contains any text
    if current_segment:
        segments.append(current_segment.strip())
    
    return segments



def translate_text(segmented_text, context_messages, max_tokens=4000):
    full_translation = ''
    
    for segment in segmented_text:
        # Trim the context to ensure it's within token limits before making a request
        while calculate_token_length(context_messages) + len(segment.split()) > max_tokens:
            context_messages.pop(0)  # Remove oldest messages if we're over the limit
        
        context_messages.append({"role": "user", "content": segment})
        try:
            response = handle_rate_limited_request(context_messages)
            translation = response.choices[0].message.content.strip()
            full_translation += translation + ' '
            # Update context_messages with the translation to maintain conversation context
            context_messages.append({"role": "assistant", "content": translation})
            # Trim context_messages if necessary
            context_messages = reduce_context_size(context_messages, max_tokens)
        except Exception as e:
            print(f"Failed to translate segment: {segment}. Error: {e}")
            break
    
    return full_translation.strip()


def format_number(text):
    # Formats numbers in text to words if they are less than 10
    new_text = ''
    for word in text.split():
        if word.isdigit():
            number = int(word)
            word = p.number_to_words(number) if number < 10 else word
        new_text += word + ' '
    return new_text.strip()

def translate_docx(file_path, source_language, target_language, output_file_path):
    doc = Document(file_path)
    new_doc = Document()
    
    context_messages = [{"role": "system", "content": f"Translate from {source_language} to {target_language}."}]
    
    for i, paragraph in enumerate(doc.paragraphs[35:]):
        if paragraph.text.strip():  # Check if the paragraph is not empty
            formatted_text = format_number(paragraph.text)
            segmented_text = segment_text(formatted_text)
            translation = translate_text(segmented_text, context_messages)
            
            trans_para = new_doc.add_paragraph(translation)
            print(f"{trans_para.text} ix-> {i}")
            new_doc.save(output_file_path)
            if paragraph.style.name.startswith('Heading'):
                trans_para.style = paragraph.style
    
    new_doc.save(output_file_path)

source_language = "Chinese"
target_language = "English"
file_path = "国史家事.docx"
output_file_path = "translated_output.docx"

translate_docx(file_path, source_language, target_language, output_file_path)

